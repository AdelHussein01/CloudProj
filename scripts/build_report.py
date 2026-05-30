from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "final"
ASSETS = ROOT / "outputs" / "report-assets"
DOCX = OUT / "cloud-native-gitops-report.docx"

TEAM_MEMBERS = [
    ("Ammar Darwish", "73226"),
    ("Djonacy Sigawuke Jr", "74294"),
    ("Adel Hussein", "73405"),
    ("Daniel Kandemiri", "73506"),
    ("Assil Sediri", "73476"),
]


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/segoeuib.ttf",
            "C:/Windows/Fonts/calibrib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
        ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str) -> None:
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    total_h = len(lines) * 24
    y = y1 + ((y2 - y1) - total_h) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font(18, True))
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, fill="#111827", font=font(18, True))
        y += 24


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#345995") -> None:
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 12, ey - 8), (ex - 12, ey + 8)]
    else:
        points = [(ex, ey), (ex + 12, ey - 8), (ex + 12, ey + 8)]
    draw.polygon(points, fill=color)


def create_architecture_diagram() -> Path:
    path = ASSETS / "architecture.png"
    img = Image.new("RGB", (1400, 760), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw.text((40, 28), "Cloud-Native GitOps Architecture", fill="#0B2545", font=font(34, True))
    draw.text((40, 72), "Git is the source of truth; ArgoCD reconciles AWS EKS to the declared state.", fill="#475569", font=font(19))

    boxes = {
        "dev": (40, 150, 230, 240),
        "github": (310, 150, 520, 240),
        "actions": (600, 150, 850, 240),
        "ghcr": (930, 130, 1160, 210),
        "helm": (930, 240, 1160, 320),
        "argo": (600, 390, 850, 500),
        "eks": (930, 390, 1280, 610),
        "secrets": (310, 430, 520, 520),
    }

    draw_box(draw, boxes["dev"], "Developer\npull request", "#DBEAFE", "#2563EB")
    draw_box(draw, boxes["github"], "GitHub repo\nAdelHussein01/CloudProj", "#DCFCE7", "#16A34A")
    draw_box(draw, boxes["actions"], "GitHub Actions\nCI + image promotion", "#FEF3C7", "#D97706")
    draw_box(draw, boxes["ghcr"], "GHCR\nsha-tagged images", "#E0E7FF", "#4F46E5")
    draw_box(draw, boxes["helm"], "Helm values\nimage tags", "#E0F2FE", "#0284C7")
    draw_box(draw, boxes["argo"], "ArgoCD\npull + self-heal", "#FCE7F3", "#DB2777")
    draw_box(draw, boxes["eks"], "AWS EKS\nNext.js web pods\nNestJS API pod\nALB ingress", "#ECFCCB", "#65A30D")
    draw_box(draw, boxes["secrets"], "AWS Secrets Manager\nExternal Secrets + IRSA", "#FEE2E2", "#DC2626")

    arrow(draw, (230, 195), (310, 195))
    arrow(draw, (520, 195), (600, 195))
    arrow(draw, (850, 170), (930, 170))
    arrow(draw, (850, 220), (930, 280))
    arrow(draw, (1045, 320), (850, 420))
    arrow(draw, (850, 445), (930, 445))
    arrow(draw, (520, 475), (930, 550))
    arrow(draw, (1045, 210), (1060, 390))

    img.save(path)
    return path


def create_pipeline_diagram() -> Path:
    path = ASSETS / "pipeline.png"
    img = Image.new("RGB", (1400, 520), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((40, 28), "GitOps CI/CD Flow", fill="#0B2545", font=font(34, True))
    steps = [
        ("1", "Pull request", "Review code and run CI"),
        ("2", "Merge to main", "Trigger release workflow"),
        ("3", "Build images", "Tag as sha-{commit}"),
        ("4", "Scan images", "Block critical/high issues"),
        ("5", "Commit values", "Promote immutable tags"),
        ("6", "ArgoCD sync", "Deploy and self-heal"),
    ]
    x = 50
    for number, title, body in steps:
        draw.rounded_rectangle((x, 150, x + 190, 350), radius=16, fill="#F8FAFC", outline="#CBD5E1", width=2)
        draw.ellipse((x + 18, 170, x + 62, 214), fill="#2563EB")
        draw.text((x + 34, 179), number, fill="#FFFFFF", font=font(18, True))
        draw.text((x + 20, 238), title, fill="#0F172A", font=font(21, True))
        draw.multiline_text((x + 20, 276), body, fill="#475569", font=font(17), spacing=4)
        if number != "6":
            arrow(draw, (x + 200, 250), (x + 245, 250), "#64748B")
        x += 220
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9 if not bold else 9.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True)
        set_cell_shading(table.rows[0].cells[index], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    doc.add_paragraph()


def add_team_table(doc: Document) -> None:
    doc.add_paragraph("Team Members").runs[0].bold = True
    add_table(
        doc,
        ["Student Name", "Student ID"],
        [[name, student_id] for name, student_id in TEAM_MEMBERS],
    )


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def build_docx() -> Path:
    ensure_dirs()
    architecture = create_architecture_diagram()
    pipeline = create_pipeline_diagram()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cloud-Native CI/CD Pipelines with GitOps")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Next.js + NestJS real-time game platform deployed with GitHub Actions, ArgoCD, and AWS EKS")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(71, 85, 105)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Repository: ").bold = True
    meta.add_run("https://github.com/AdelHussein01/CloudProj")
    add_team_table(doc)

    add_heading(doc, "Executive Summary")
    doc.add_paragraph(
        "This project implements a GitOps-based CI/CD pipeline for a cloud-native two-player game platform. "
        "The application lets a host create a private link, choose XO or rock-paper-scissors, and invite the first user who opens the link to join the selected game in real time."
    )
    doc.add_paragraph(
        "The technical goal is to show how declarative infrastructure, immutable image promotion, automated testing, secure secret management, rollback, and observability combine into a reliable delivery workflow."
    )

    add_heading(doc, "Project Scope")
    add_bullets(
        doc,
        [
            "Application: Next.js frontend and NestJS Socket.IO API.",
            "CI: GitHub Actions for type checks, tests, builds, Docker image publishing, and image scanning.",
            "CD: ArgoCD watches Helm manifests and reconciles AWS EKS to the Git-declared desired state.",
            "Secrets: AWS Secrets Manager, External Secrets Operator, and IAM Roles for Service Accounts.",
            "Rollback: Git revert or a controlled rollback workflow that commits known-good image tags.",
            "AWS execution is intentionally left as the final phase to protect account access and control cost.",
        ],
    )

    add_heading(doc, "System Architecture")
    doc.add_picture(str(architecture), width=Inches(6.5))
    doc.add_paragraph(
        "The architecture separates CI from CD. GitHub Actions validates code and promotes immutable image tags to Git, while ArgoCD performs cluster reconciliation from inside Kubernetes."
    )

    add_heading(doc, "Application Design")
    add_table(
        doc,
        ["Layer", "Technology", "Responsibility"],
        [
            ["Web", "Next.js", "Room creation, share links, join screen, XO board, RPS rounds, responsive UI."],
            ["API", "NestJS + Socket.IO", "Real-time rooms, player state, game validation, broadcasts, health endpoint."],
            ["Packaging", "Docker", "Separate production images for web and API."],
            ["Runtime", "Kubernetes", "Deployments, services, ingress, probes, resources, and optional autoscaling."],
        ],
    )
    doc.add_paragraph(
        "The API stores room state in memory and therefore starts with one API replica for the demo. A production version should add Redis or another shared state backend before horizontal API scaling."
    )

    add_heading(doc, "CI/CD Pipeline")
    doc.add_picture(str(pipeline), width=Inches(6.5))
    add_numbered(
        doc,
        [
            "A developer opens a pull request.",
            "GitHub Actions runs type checks, tests, and production builds.",
            "After merge to main, Docker images are built and pushed to GitHub Container Registry.",
            "Images are tagged with the commit SHA to preserve immutability.",
            "Trivy scans the images and blocks high or critical vulnerabilities.",
            "The release workflow commits updated Helm image tags.",
            "ArgoCD detects the Git change and syncs AWS EKS.",
        ],
    )

    add_heading(doc, "GitOps vs Traditional CI/CD")
    add_table(
        doc,
        ["Area", "Traditional CI/CD", "GitOps in This Project"],
        [
            ["Deployment actor", "Pipeline pushes to the cluster.", "ArgoCD pulls desired state from Git."],
            ["Source of truth", "Pipeline state and live cluster can diverge.", "Git stores the declared runtime state."],
            ["Rollback", "Re-run a pipeline or manually deploy an older artifact.", "Revert Git or commit previous immutable tags."],
            ["Drift handling", "Manual checks or periodic audits.", "ArgoCD detects and self-heals drift."],
            ["Audit trail", "Pipeline logs plus release notes.", "Git history plus ArgoCD sync history."],
        ],
    )

    add_heading(doc, "Security and Secret Management")
    add_bullets(
        doc,
        [
            "No raw secrets are committed to Git.",
            "GitHub Actions uses scoped permissions for CI, package publishing, and rollback.",
            "Runtime secrets are stored in AWS Secrets Manager and synced by External Secrets Operator.",
            "External Secrets authenticates through IRSA, avoiding static AWS keys in pods.",
            "Image tags use commit SHAs, not mutable production tags.",
            "Container images are scanned before promotion.",
        ],
    )

    add_heading(doc, "Rollback Strategy")
    doc.add_paragraph(
        "The preferred rollback is a Git operation. Reverting the GitOps promotion commit or running the manual rollback workflow updates Helm values with known-good image tags. ArgoCD then reconciles the cluster to that older declared state."
    )
    add_bullets(
        doc,
        [
            "Fast rollback path: run the rollback workflow with previous `sha-*` image tags.",
            "Audit-friendly rollback path: revert the bad promotion commit.",
            "Avoid cluster-only rollback because it creates drift between Git and Kubernetes.",
        ],
    )

    add_heading(doc, "Observability and Evaluation")
    add_table(
        doc,
        ["Metric", "Purpose", "How to Measure"],
        [
            ["Deployment frequency", "Shows release throughput.", "Count successful release workflow runs."],
            ["Lead time", "Measures merge-to-production speed.", "Time from merge to healthy ArgoCD sync."],
            ["Change failure rate", "Measures deployment quality.", "Failed deployments divided by total deployments."],
            ["MTTR", "Measures recovery speed.", "Time from bad release detection to healthy rollback."],
            ["Drift repair time", "Measures GitOps self-healing.", "Time from manual cluster drift to ArgoCD correction."],
        ],
    )

    add_heading(doc, "Demo Plan")
    add_numbered(
        doc,
        [
            "Open the deployed application.",
            "Create an XO room and copy the share link.",
            "Join from a second browser and finish an XO match.",
            "Create a rock-paper-scissors room and play multiple rounds.",
            "Open GitHub Actions and show CI and release workflows.",
            "Open ArgoCD and show sync status and application health.",
            "Create controlled drift by scaling the web deployment and show ArgoCD self-healing.",
            "Run the rollback workflow or revert a promotion commit and verify the older version is restored.",
        ],
    )

    add_heading(doc, "AWS Final Phase")
    doc.add_paragraph(
        "AWS deployment should be performed last. This avoids exposing credentials early, gives time to review the project artifacts, and allows the team to confirm cost and cleanup expectations before creating paid cloud resources."
    )
    add_bullets(
        doc,
        [
            "Create EKS infrastructure with Terraform.",
            "Install AWS Load Balancer Controller, External Secrets Operator, and ArgoCD.",
            "Create the required AWS Secrets Manager entry.",
            "Apply ArgoCD manifests and verify the application sync.",
            "Use a real domain or the AWS ALB DNS name for the demo.",
        ],
    )

    add_heading(doc, "Conclusion")
    doc.add_paragraph(
        "GitOps improves reliability by making Git the system of record for deployment state. In this project, every deployment and rollback is traceable, automated, and reproducible. The main engineering risks are cloud IAM setup, secret handling, and state sharing for real-time WebSocket workloads."
    )

    add_heading(doc, "References")
    add_bullets(
        doc,
        [
            "Amazon EKS User Guide: https://docs.aws.amazon.com/eks/latest/userguide/",
            "AWS Load Balancer Controller on EKS: https://docs.aws.amazon.com/eks/latest/userguide/lbc-helm.html",
            "ArgoCD documentation: https://argo-cd.readthedocs.io/",
            "GitHub Actions documentation: https://docs.github.com/en/actions",
            "GitHub Container Registry documentation: https://docs.github.com/en/packages/working-with-a-github-packages-registry/working-with-the-container-registry",
            "External Secrets Operator documentation: https://external-secrets.io/",
        ],
    )

    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    print(build_docx())
